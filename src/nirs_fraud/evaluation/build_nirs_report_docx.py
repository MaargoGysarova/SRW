from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from ..catalog.taxonomy import SCENARIOS, SIGNALS


ROOT = Path(__file__).resolve().parents[3]
OUTPUT_DIR = ROOT / "outputs"
REPORT_DIR = OUTPUT_DIR / "nirs_report"
REPORT_PATH = REPORT_DIR / "nirs_report_gusarova_2026.docx"

STUDENT_NAME = "М.О. Гусарова"
GROUP = "6512-100503D"
SUPERVISOR = "к.т.н., доцент Митекин В.А."
PRACTICE_START = "01.03.2026"
PRACTICE_END = "10.06.2026"
SUBMIT_DATE = "10.06.2026"
DEFENSE_DATE = "10.06.2026"
UNIVERSITY = (
    "федеральное государственное автономное образовательное учреждение высшего образования "
    "«Самарский национальный исследовательский университет имени академика С.П. Королева»"
)
DEPARTMENT = "Кафедра геоинформатики и информационной безопасности"
FACULTY = "Институт информатики и кибернетики"
TOPIC = (
    "Исследование методов выявления мошеннических диалогов на русском языке "
    "с использованием правиловых и LLM-подходов"
)

ARCHITECTURE_NAMES_RU = {
    "single_llm": "конфигурация прямой LLM-классификации",
    "llm_checklist": "конфигурация с чек-листом признаков",
    "llm_self_check": "конфигурация с самопроверкой",
    "llm_ensemble": "ансамблевая конфигурация",
}


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def to_float(value: str) -> float:
    return float(value)


def fmt(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size: int = 11) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 14), ("Heading 3", 13)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    bold: bool = False,
    italic: bool = False,
    size: int = 14,
    first_line_indent: float | None = 1.25,
    style: str | None = None,
) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = align
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_code_block(document: Document, lines: list[str], *, size: int = 11) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(1.0)
        paragraph.paragraph_format.right_indent = Cm(1.0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_paragraph(style=f"Heading {level}")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.first_line_indent = Cm(0)
    run = heading.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level < 3 else 13)
    run.bold = True


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_number_item(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].width = Cm(widths_cm[index])
        set_cell_text(header_cells[index], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            row.cells[index].width = Cm(widths_cm[index])
            align = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[index], value, align=align)

    add_paragraph(document, "", first_line_indent=None)


def add_caption(document: Document, text: str) -> None:
    add_paragraph(
        document,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=False,
        italic=False,
        size=12,
        first_line_indent=None,
    )


def add_figure(document: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(document, caption)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def add_footer_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def metric_map(rows: list[dict[str, str]], key: str) -> dict[str, dict[str, str]]:
    return {row[key]: row for row in rows}


def pair_map(rows: list[dict[str, str]], key1: str, key2: str) -> dict[tuple[str, str], dict[str, str]]:
    return {(row[key1], row[key2]): row for row in rows}


def build_report() -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    exp1_rules = load_csv(OUTPUT_DIR / "experiment_01_2" / "metrics.csv")[0]
    exp1_qwen = metric_map(
        load_csv(OUTPUT_DIR / "experiment_01_2" / "local_llm_experiment_01_Qwen_Qwen2_5-14B-Instruct_metrics.csv"),
        "architecture",
    )
    exp1_old = metric_map(
        load_csv(OUTPUT_DIR / "experiment_01" / "local_llm_experiment_01_Qwen_Qwen2_5-14B-Instruct_metrics.csv"),
        "architecture",
    )
    exp2_rules = metric_map(load_csv(OUTPUT_DIR / "experiment_02_2" / "experiment_02_metrics.csv"), "variant")
    exp2_qwen = pair_map(
        load_csv(OUTPUT_DIR / "experiment_02_2" / "local_llm_experiment_02_Qwen_Qwen2_5-14B-Instruct_metrics.csv"),
        "architecture",
        "variant",
    )
    exp3_rules = load_csv(OUTPUT_DIR / "experiment_03" / "experiment_03_metrics.csv")[0]
    exp3_qwen = metric_map(
        load_csv(OUTPUT_DIR / "experiment_03" / "local_llm_experiment_03_Qwen_Qwen2_5-14B-Instruct_metrics.csv"),
        "architecture",
    )
    gemma_exp1 = load_csv(OUTPUT_DIR / "experiment_04" / "experiment_01" / "gemma_experiment_01_gemma-4-31b-it_metrics.csv")[0]

    exp1_predictions = load_csv(OUTPUT_DIR / "experiment_01_2" / "predictions.csv")
    exp2_predictions = load_csv(OUTPUT_DIR / "experiment_02_2" / "experiment_02_predictions.csv")
    exp3_predictions = load_csv(OUTPUT_DIR / "experiment_03" / "experiment_03_predictions.csv")

    exp1_size = len(exp1_predictions)
    exp2_size = len(exp2_predictions)
    exp3_size = len(exp3_predictions)

    document = Document()
    configure_styles(document)
    add_footer_page_number(document.sections[0])

    # Title page
    add_paragraph(document, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, first_line_indent=None)
    add_paragraph(document, UNIVERSITY, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    add_paragraph(document, "(Самарский университет)", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, FACULTY, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    add_paragraph(document, DEPARTMENT, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, "ОТЧЕТ ПО ПРАКТИКЕ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line_indent=None)
    add_paragraph(document, "по научно-исследовательской работе", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, f"Тема работы: «{TOPIC}»", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, "Вид практики: производственная практика", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(document, "Тип практики: научно-исследовательская работа", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(
        document,
        f"Сроки прохождения практики: с {PRACTICE_START} г. по {PRACTICE_END} г.",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
    )
    add_paragraph(
        document,
        "по направлению подготовки 10.05.03 Информационная безопасность автоматизированных систем",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
    )
    add_paragraph(
        document,
        "направленность (профиль) «Безопасность открытых информационных систем»",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
    )
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, f"Обучающийся группы № {GROUP} {STUDENT_NAME}", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(document, f"Руководитель практики от университета {SUPERVISOR}", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, f"Дата сдачи: {SUBMIT_DATE} г.", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(document, f"Дата защиты: {DEFENSE_DATE} г.", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(document, "Оценка ___________________", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, "Самара 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

    # Contents
    add_page_break(document)
    add_heading(document, "СОДЕРЖАНИЕ", 1)
    for item in [
        "ВВЕДЕНИЕ",
        "ВЫПОЛНЕНИЕ ЗАДАНИЯ",
        "1.1 Актуальность задачи выявления мошеннических диалогов",
        "1.2 Постановка задачи классификации и структура признаков",
        "1.3 Формирование набора данных и экспериментального стенда",
        "1.4 Исследуемые архитектуры и методика оценки качества",
        "1.5 Результаты базового эксперимента",
        "1.6 Результаты эксперимента устойчивости к аугментациям",
        "1.7 Результаты внешнего бенчмарка",
        "1.8 Дополнительные исследования",
        "1.9 Анализ полученных результатов",
        "1.10 Выводы по выполненному заданию",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "ОТЗЫВ О ПРОХОЖДЕНИИ ПРАКТИКИ",
    ]:
        add_paragraph(document, item, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)

    # Individual assignment
    add_page_break(document)
    add_heading(document, "Индивидуальное задание на практику", 1)
    add_paragraph(document, f"Студенту группы № {GROUP} {STUDENT_NAME}", first_line_indent=None)
    add_paragraph(
        document,
        "Направление на практику оформлено приказом по университету от 26.02.2026 № 119-ПР "
        "на кафедру геоинформатики и информационной безопасности Самарского университета.",
        first_line_indent=None,
    )
    competencies = [
        (
            "ОПК-10",
            "Знать, уметь и владеть методами защиты информации при решении профессиональных задач.",
            "Изучены и применены методы выявления мошеннических коммуникаций, включая сигнатурное выделение опасных паттернов, "
            "анализ сценариев социальной инженерии и классификацию диалогов по уровню риска.",
        ),
        (
            "ОПК-11",
            "Знать методы разработки компонентов систем защиты информации автоматизированных систем.",
            "Разработаны программные компоненты экспериментального стенда: модули подготовки датасета, генерации и валидации синтетических диалогов, "
            "шаблоны промптов, скрипты запуска моделей и расчета метрик.",
        ),
        (
            "ОПК-2",
            "Использовать базовые понятия информатики для решения профессиональных задач.",
            "Применены базовые подходы анализа данных и машинного обучения: разметка классов, структурирование JSON-ответов, "
            "агрегация прогнозов и воспроизводимое проведение вычислительных экспериментов.",
        ),
        (
            "ОПК-3.1",
            "Использовать математический аппарат и алгоритмические методы в профессиональной деятельности.",
            "Использованы алгоритмы классификации на основе правил, пошаговой самопроверки, ансамблирования и постобработки результатов, "
            "а также методы сравнения нескольких архитектур по единому набору показателей качества.",
        ),
        (
            "ОПК-3.5",
            "Использовать подходы и методы теории вероятности при анализе прикладных задач.",
            "Проведен вероятностный и статистический анализ качества классификации с использованием метрик accuracy, precision, recall и F1, "
            "а также анализа ложноположительных и ложноотрицательных ошибок.",
        ),
        (
            "ОПК-4.2",
            "Использовать модели и количественные методы для решения научно-технических задач в области информационной безопасности.",
            "Разработана методика количественной оценки устойчивости детектора к перефразированию, завуалированию признаков мошенничества и ASR-подобному шуму.",
        ),
        (
            "ОПК-5.1.1",
            "Использовать понятийный аппарат открытых информационных систем и их взаимодействия с внешними системами.",
            "Исследуемая задача рассмотрена в контексте открытых информационных систем, где мошеннический диалог выступает интерфейсом атакующего воздействия "
            "на пользователя через телефонные и онлайн-каналы.",
        ),
        (
            "ОПК-7.1",
            "Использовать инструментальные средства разработки и отладки программного обеспечения.",
            "В работе использовались Python, JSONL/CSV-пайплайны, локальные и API-модели, а также скрипты визуализации результатов и автоматического формирования отчетных материалов.",
        ),
        (
            "ОПК-8.1",
            "Применять методы и программные средства автоматизации научных исследований.",
            "Автоматизированы подготовка промптов, запуск экспериментальных серий, сохранение частичных результатов, расчет метрик, генерация таблиц и построение графиков.",
        ),
        (
            "ОПК-8.2",
            "Пользоваться научно-технической литературой и источниками цитирования.",
            "Проведен обзор научных публикаций по LLM, scam detection, prompt engineering и text augmentation; подготовлен список использованных источников для итоговой записки.",
        ),
        (
            "ОПК-9.1",
            "Разрабатывать решения с применением знаний теории информации и методов кодирования данных.",
            "При анализе диалогов использованы представления текстовых сообщений как носителей информативных сигналов мошенничества и семантических индикаторов риска.",
        ),
        (
            "ОПК-9.5",
            "Применять современные программные, технические средства и информационные технологии в системах передачи информации.",
            "Использованы современные LLM-инструменты, локальные inference-пайплайны и API-доступ к облачным моделям для сравнительной оценки методов классификации.",
        ),
    ]
    add_table(
        document,
        ["Компетенция", "Планируемые результаты практики", "Содержание задания"],
        competencies,
        [2.5, 6.0, 8.0],
    )
    add_paragraph(document, f"Дата выдачи задания: {PRACTICE_START} г.", first_line_indent=None)
    add_paragraph(document, f"Срок предоставления отчета на кафедру: {PRACTICE_END} г.", first_line_indent=None)
    add_paragraph(document, f"Руководитель практики от университета: {SUPERVISOR}", first_line_indent=None)
    add_paragraph(document, "Задание принял к исполнению ____________________", first_line_indent=None)

    # Main report
    add_page_break(document)
    add_heading(document, "О Т Ч Е Т", 1)
    add_paragraph(document, "о выполнении индивидуального задания по научно-исследовательской работе", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, first_line_indent=None)

    add_heading(document, "ВВЕДЕНИЕ", 1)
    add_paragraph(
        document,
        "В ходе прохождения практики по научно-исследовательской работе была поставлена задача разработать и экспериментально "
        "оценить методы выявления мошеннических диалогов на русском языке. Практическая значимость этой задачи определяется тем, "
        "что значительная часть современных атак на граждан и сотрудников организаций реализуется не через эксплуатацию программной "
        "уязвимости, а через социальную инженерию: злоумышленник убеждает жертву сообщить код подтверждения, установить приложение "
        "удаленного доступа, перейти по фишинговой ссылке или самостоятельно перевести деньги на «безопасный счет».",
    )
    add_paragraph(
        document,
        "Цель исследования состояла в сравнении эффективности базового классификатора на основе правил и нескольких LLM-конфигураций для "
        "классификации коротких диалогов по трем классам: fraud, suspicious и safe. Для достижения цели были поставлены следующие задачи: "
        "сформировать каталог сценариев и сигналов мошенничества, подготовить базовый датасет и его аугментированные варианты, "
        "реализовать правила и LLM-архитектуры, определить единый протокол оценки качества и провести серию сравнительных экспериментов.",
    )
    add_paragraph(
        document,
        "Объектом исследования являются русскоязычные диалоги, содержащие признаки банковского, телефонного и платформенного мошенничества. "
        "Предметом исследования являются методы автоматической классификации таких диалогов, включая детерминированные правила, "
        "одиночные LLM-запросы, чек-листовые промпты, самопроверяющиеся схемы и ансамбли ответов.",
    )
    add_paragraph(
        document,
        "В работе были последовательно реализованы четыре направления: построение корпуса и таксономии, разработка экспериментального "
        "стенда, оценка качества на базовом наборе и анализ устойчивости моделей к переформулировкам, завуалированным сигналам и "
        "ASR-подобным искажениям текста.",
    )

    add_heading(document, "ВЫПОЛНЕНИЕ ЗАДАНИЯ", 1)

    add_heading(document, "1.1 Актуальность задачи выявления мошеннических диалогов", 2)
    add_paragraph(
        document,
        "Телефонное и мессенджерное мошенничество остается одной из наиболее опасных форм атак на пользователей информационных систем. "
        "В отличие от классических технических атак, социальная инженерия строится на манипуляции доверием, срочностью и авторитетом. "
        "Вредоносное воздействие происходит на уровне естественного языка, поэтому для автоматического выявления угроз требуется анализ "
        "не только отдельных ключевых слов, но и общего сценария разговора, роли собеседников и контекста требуемого действия.",
    )
    add_paragraph(
        document,
        "Дополнительную сложность создают несколько факторов. Во-первых, мошеннические схемы постоянно варьируются и адаптируются под "
        "информационную повестку: звонки от имени банка, следователя, оператора связи, Госуслуг, службы доставки или маркетплейса "
        "могут использовать сходные психологические приемы при различающейся лексике. Во-вторых, опасные инструкции часто выражаются "
        "непрямо: вместо прямой просьбы сообщить код из СМС злоумышленник может говорить о «подтверждении операции» или «цифрах из уведомления». "
        "В-третьих, на практике нередко встречаются ASR-подобные искажения, речевой шум и краткие фрагменты диалога, что усложняет "
        "задачу для традиционных сигнатурных подходов.",
    )
    add_paragraph(
        document,
        "По этой причине исследование LLM-подходов к анализу мошеннических диалогов представляется актуальным. Большие языковые модели "
        "способны учитывать семантику запроса, скрытые манипулятивные паттерны и связь между отдельными репликами. При этом для практического "
        "использования необходимо проверить, насколько такие модели действительно превосходят простые правила, насколько устойчивы к "
        "переформулировкам и можно ли рассчитывать на переносимость результатов на внешние данные.",
    )

    add_heading(document, "1.2 Постановка задачи классификации и структура признаков", 2)
    add_paragraph(
        document,
        "В работе решалась задача трехклассовой классификации диалогов. Класс fraud присваивался в случаях, когда в тексте присутствовало "
        "явное мошенническое действие или социально-инженерный сценарий: выдача за банк или госорган, давление по времени, просьба назвать "
        "код из СМС, установить приложение, перейти по ссылке, перевести деньги либо скрыть разговор от окружающих. Класс suspicious "
        "использовался для пограничных случаев, содержащих тревожные признаки, но не дающих полной уверенности в наличии фрода. Класс safe "
        "соответствовал бытовым и служебным безопасным диалогам без чувствительных действий.",
    )
    add_paragraph(
        document,
        f"Для обеспечения содержательной полноты был сформирован каталог из {len(SCENARIOS)} сценариев. Он включает банковые, "
        "правоохранительные, семейно-эмоциональные, маркетплейсные, инвестиционные и иные типы мошеннических коммуникаций, а также "
        "безопасные контрольные сценарии. Одновременно был задан набор из "
        f"{len(SIGNALS)} канонических сигналов мошенничества, среди которых urgency, pressure, bank_impersonation, sms_code_request, "
        "safe_account_transfer, secrecy, remote_access_app и phishing_link. Такая декомпозиция позволила фиксировать не только итоговую метку, "
        "но и интерпретируемые причины принятого решения.",
    )
    add_paragraph(
        document,
        "Каждый экспериментальный пример представлялся в виде короткого русскоязычного диалога, для которого модель должна была вернуть "
        "JSON-объект со значениями label, fraud_score, signals и explanation. Тем самым оценивалась не только способность выбрать класс, "
        "но и умение модели выделять внутренние семантические признаки мошеннического сценария.",
    )
    add_paragraph(
        document,
        "Пример ожидаемой структуры ответа модели приведен ниже.",
    )
    add_code_block(
        document,
        [
            "{",
            '  "label": "fraud",',
            '  "fraud_score": 0.97,',
            '  "signals": ["urgency", "sms_code_request", "bank_impersonation"],',
            '  "explanation": "Собеседник представляется сотрудником банка, оказывает давление по времени и просит сообщить код из СМС."', 
            "}",
        ],
    )
    add_caption(document, "Пример 1 – Ожидаемая JSON-структура ответа модели")

    add_heading(document, "1.3 Формирование набора данных и экспериментального стенда", 2)
    add_paragraph(
        document,
        "Формирование набора данных осуществлялось в три последовательных этапа: генерация, аугментация и валидация. "
        "На первом этапе с помощью LLM-генератора создавались синтетические русскоязычные диалоги, соответствующие заданным сценариям, "
        "целевым меткам и наборам канонических сигналов. На втором этапе часть примеров расширялась с помощью аугментатора, "
        "формировавшего дополнительные варианты диалогов с перефразированием, менее явным выражением мошеннических признаков "
        "и ASR-подобными искажениями текста. На третьем этапе полученные записи проходили LLM-валидацию на соответствие метке, "
        "сценарию, набору сигналов и требованиям к качеству датасета.\n\n"
        f"Итоговый базовый набор Эксперимента 1 содержал {exp1_size} диалогов: 35 примеров класса fraud, 20 примеров класса suspicious "
        "и 30 примеров класса safe. Он использовался как основная контрольная выборка для сравнения базового классификатора на основе правил "
        "и четырех LLM-конфигураций. Для исследования устойчивости была подготовлена расширенная выборка Эксперимента 2 объемом "
        f"{exp2_size} примера, включавшая исходные записи и три типа аугментаций: paraphrase, subtle и asr_noise. "
        "Внешний бенчмарк Эксперимента 3 был сформирован на основе поднабора открытого датасета Scam and Non-Scam Call Conversation Dataset, "
        "опубликованного на платформе Kaggle (https://www.kaggle.com/datasets/teeconnie/scam-and-non-scam-call-conversation-dataset/data). "
        f"Итоговая выборка содержала {exp3_size} диалог, из которых 30 относились к классу fraud, а 31 — к классу safe, "
        "что позволило оценить переносимость предложенных подходов вне базовой синтетической выборки.",
    )
    add_paragraph(
        document,
        "Экспериментальный стенд был реализован как набор воспроизводимых Python-скриптов. Он включал каталог сценариев и сигналов, "
        "шаблоны промптов для генерации, аугментации, валидации и классификации, средства нормализации ответов модели, подсчет метрик, "
        "формирование CSV-отчетов и построение графиков. Такая организация позволила не только провести единичный запуск, но и "
        "многократно переиспользовать пайплайн при изменении конфигурации промптов или модели.",
    )
    add_paragraph(
        document,
        "Особое внимание было уделено контролю формата. Для LLM-запросов были подготовлены строгие JSON-инструкции, а все ответы "
        "проходили постобработку и нормализацию. Благодаря этому удалось сравнивать архитектуры в единых условиях и минимизировать "
        "влияние ошибок сериализации на итоговые метрики.",
    )
    add_table(
        document,
        ["Набор", "Размер", "Распределение классов", "Назначение"],
        [
            ["Эксперимент 1", str(exp1_size), "fraud=35, suspicious=20, safe=30", "Сравнение базового классификатора на основе правил и LLM"],
            ["Эксперимент 2", str(exp2_size), "fraud=137, suspicious=77, safe=30", "Оценка устойчивости к аугментациям"],
            ["Эксперимент 3", str(exp3_size), "fraud=30, safe=31", "Проверка переносимости на внешний набор данных"],
        ],
        [3.0, 2.5, 5.0, 5.0],
    )
    add_caption(document, "Таблица 1 – Характеристика наборов данных, использованных в экспериментах")

    add_heading(document, "1.4 Исследуемые архитектуры и методика оценки качества", 2)
    add_paragraph(
        document,
        "В исследовании сравнивались пять архитектур. Первая архитектура представляла собой базовый классификатор на основе правил, "
        "построенный на детерминированных эвристиках по ключевым словам и опасным паттернам. Вторая архитектура использовала "
        "прямой одиночный запрос к модели. Третья архитектура задавала модели явный чек-лист признаков социальной инженерии. "
        "Четвертая архитектура строилась по схеме первичного ответа с последующей самопроверкой и возможной корректировкой вывода. "
        "Пятая архитектура моделировала агрегирование нескольких независимых точек зрения на один и тот же диалог.",
    )
    add_paragraph(
        document,
        "В качестве основных метрик использовались accuracy, precision_fraud, recall_fraud и f1_fraud. Выбор метрик был обусловлен "
        "прикладным характером задачи: для систем обнаружения мошенничества особенно важно уметь не пропускать опасные случаи, поэтому "
        "метрика recall для класса fraud играет ключевую роль. В то же время рост ложноположительных срабатываний ухудшает практическую "
        "полезность детектора, поэтому дополнительно анализировались false positives и false negatives. В каждом эксперименте также "
        "фиксировалась средняя задержка обработки одного примера.",
    )
    add_paragraph(
        document,
        "Программа исследования включала три основных эксперимента и два дополнительных. В первом эксперименте сравнивались все "
        "архитектуры на базовом синтетическом наборе. Во втором исследовалась устойчивость моделей к переформулировкам, неявному "
        "выражению мошеннических признаков и ASR-подобным искажениям текста. В третьем эксперименте проверялась переносимость "
        "подходов на внешний набор данных. Дополнительно были проанализированы две версии промптовой конфигурации первого эксперимента "
        "и отдельный API-прогон модели Gemma.",
    )

    add_heading(document, "1.5 Результаты базового эксперимента", 2)
    exp1_rows = [
        ["Правила", fmt(to_float(exp1_rules["accuracy"])), fmt(to_float(exp1_rules["precision_fraud"])), fmt(to_float(exp1_rules["recall_fraud"])), fmt(to_float(exp1_rules["f1_fraud"])), exp1_rules["false_positives"], exp1_rules["false_negatives"]],
        ["Qwen 2.5-14B: прямая классификация", fmt(to_float(exp1_qwen["single_llm"]["accuracy"])), fmt(to_float(exp1_qwen["single_llm"]["precision_fraud"])), fmt(to_float(exp1_qwen["single_llm"]["recall_fraud"])), fmt(to_float(exp1_qwen["single_llm"]["f1_fraud"])), exp1_qwen["single_llm"]["false_positives"], exp1_qwen["single_llm"]["false_negatives"]],
        ["Qwen 2.5-14B: чек-лист признаков", fmt(to_float(exp1_qwen["llm_checklist"]["accuracy"])), fmt(to_float(exp1_qwen["llm_checklist"]["precision_fraud"])), fmt(to_float(exp1_qwen["llm_checklist"]["recall_fraud"])), fmt(to_float(exp1_qwen["llm_checklist"]["f1_fraud"])), exp1_qwen["llm_checklist"]["false_positives"], exp1_qwen["llm_checklist"]["false_negatives"]],
        ["Qwen 2.5-14B: самопроверка", fmt(to_float(exp1_qwen["llm_self_check"]["accuracy"])), fmt(to_float(exp1_qwen["llm_self_check"]["precision_fraud"])), fmt(to_float(exp1_qwen["llm_self_check"]["recall_fraud"])), fmt(to_float(exp1_qwen["llm_self_check"]["f1_fraud"])), exp1_qwen["llm_self_check"]["false_positives"], exp1_qwen["llm_self_check"]["false_negatives"]],
        ["Qwen 2.5-14B: ансамбль", fmt(to_float(exp1_qwen["llm_ensemble"]["accuracy"])), fmt(to_float(exp1_qwen["llm_ensemble"]["precision_fraud"])), fmt(to_float(exp1_qwen["llm_ensemble"]["recall_fraud"])), fmt(to_float(exp1_qwen["llm_ensemble"]["f1_fraud"])), exp1_qwen["llm_ensemble"]["false_positives"], exp1_qwen["llm_ensemble"]["false_negatives"]],
    ]
    add_table(
        document,
        ["Подход", "Accuracy", "Precision", "Recall", "F1", "FP", "FN"],
        exp1_rows,
        [5.7, 1.7, 1.7, 1.7, 1.4, 1.2, 1.2],
    )
    add_caption(document, "Таблица 2 – Результаты Эксперимента 1 на базовом синтетическом наборе")
    add_paragraph(
        document,
        f"Результаты базового эксперимента показывают уверенное превосходство LLM-подходов над базовым классификатором на основе правил. "
        f"Базовый классификатор на основе правил достиг F1 fraud = {fmt(to_float(exp1_rules['f1_fraud']))} при recall fraud = {fmt(to_float(exp1_rules['recall_fraud']))}, "
        f"что указывает на заметное число пропущенных опасных случаев (FN = {exp1_rules['false_negatives']}). "
        f"Наилучший итоговый F1 продемонстрировала конфигурация прямой LLM-классификации: {fmt(to_float(exp1_qwen['single_llm']['f1_fraud']))}. "
        f"Конфигурация с самопроверкой показала максимальную полноту {fmt(to_float(exp1_qwen['llm_self_check']['recall_fraud']))}, "
        "что особенно важно для детекторов мошенничества, однако заплатила за это большим числом ложноположительных срабатываний.",
    )
    add_paragraph(
        document,
        "Полученные результаты позволяют сделать два важных вывода. Во-первых, даже одиночный запрос к LLM в данной постановке "
        "существенно превосходит набор ручных правил. Во-вторых, разные промптовые архитектуры решают задачу по-разному: "
        "прямая LLM-классификация обеспечивает лучший баланс precision и recall, конфигурация с самопроверкой повышает чувствительность "
        "к слабым сигналам, а конфигурации с чек-листом и ансамблированием обеспечивают более интерпретируемый, но не максимальный итоговый F1.",
    )
    add_figure(
        document,
        OUTPUT_DIR / "report_charts" / "exp1_f1_comparison.png",
        "Рисунок 1 – Сравнение базового классификатора на основе правил и LLM-конфигураций по F1 fraud на базовом наборе",
    )

    add_heading(document, "1.6 Результаты эксперимента устойчивости к аугментациям", 2)
    exp2_agg_rows = [
        ["прямая LLM-классификация", fmt(to_float(exp2_qwen[("single_llm", "all_augmented")]["f1_fraud"])), fmt(to_float(exp2_qwen[("single_llm", "all_variants")]["f1_fraud"])), fmt(to_float(exp2_qwen[("single_llm", "all_augmented")]["recall_fraud"])), fmt(to_float(exp2_qwen[("single_llm", "all_variants")]["recall_fraud"]))],
        ["чек-лист признаков", fmt(to_float(exp2_qwen[("llm_checklist", "all_augmented")]["f1_fraud"])), fmt(to_float(exp2_qwen[("llm_checklist", "all_variants")]["f1_fraud"])), fmt(to_float(exp2_qwen[("llm_checklist", "all_augmented")]["recall_fraud"])), fmt(to_float(exp2_qwen[("llm_checklist", "all_variants")]["recall_fraud"]))],
        ["самопроверка", fmt(to_float(exp2_qwen[("llm_self_check", "all_augmented")]["f1_fraud"])), fmt(to_float(exp2_qwen[("llm_self_check", "all_variants")]["f1_fraud"])), fmt(to_float(exp2_qwen[("llm_self_check", "all_augmented")]["recall_fraud"])), fmt(to_float(exp2_qwen[("llm_self_check", "all_variants")]["recall_fraud"]))],
        ["ансамбль", fmt(to_float(exp2_qwen[("llm_ensemble", "all_augmented")]["f1_fraud"])), fmt(to_float(exp2_qwen[("llm_ensemble", "all_variants")]["f1_fraud"])), fmt(to_float(exp2_qwen[("llm_ensemble", "all_augmented")]["recall_fraud"])), fmt(to_float(exp2_qwen[("llm_ensemble", "all_variants")]["recall_fraud"]))],
    ]
    add_table(
        document,
        ["Архитектура", "F1 all_augmented", "F1 all_variants", "Recall all_augmented", "Recall all_variants"],
        exp2_agg_rows,
        [4.4, 3.0, 3.0, 3.0, 3.0],
    )
    add_caption(document, "Таблица 3 – Агрегированные результаты Эксперимента 2 по аугментированным поднаборам")
    exp2_best_rows = [
        ["original", "прямая LLM-классификация", fmt(to_float(exp2_qwen[("single_llm", "original")]["f1_fraud"]))],
        ["paraphrase", "чек-лист признаков", fmt(to_float(exp2_qwen[("llm_checklist", "paraphrase")]["f1_fraud"]))],
        ["subtle", "самопроверка", fmt(to_float(exp2_qwen[("llm_self_check", "subtle")]["f1_fraud"]))],
        ["asr_noise", "прямая LLM-классификация", fmt(to_float(exp2_qwen[("single_llm", "asr_noise")]["f1_fraud"]))],
        ["all_augmented", "прямая LLM-классификация", fmt(to_float(exp2_qwen[("single_llm", "all_augmented")]["f1_fraud"]))],
        ["all_variants", "прямая LLM-классификация", fmt(to_float(exp2_qwen[("single_llm", "all_variants")]["f1_fraud"]))],
    ]
    add_table(
        document,
        ["Вариант", "Лучшая конфигурация", "F1 fraud"],
        exp2_best_rows,
        [4.0, 5.5, 3.0],
    )
    add_caption(document, "Таблица 4 – Лучшая конфигурация для каждого варианта аугментации")
    add_paragraph(
        document,
        f"Второй эксперимент показал, что базовый классификатор на основе правил крайне чувствителен к модификациям текста. "
        f"Для варианта asr_noise его F1 fraud снизился до {fmt(to_float(exp2_rules['asr_noise']['f1_fraud']))}, "
        f"тогда как лучший LLM-результат на этом же поднаборе составил {fmt(to_float(exp2_qwen[('single_llm', 'asr_noise')]['f1_fraud']))}. "
        f"На агрегированном наборе all_augmented лучшей конфигурацией стала прямая LLM-классификация с F1 fraud = "
        f"{fmt(to_float(exp2_qwen[('single_llm', 'all_augmented')]['f1_fraud']))}, а на самом сложном subtle-поднаборе "
        f"лидировала конфигурация с самопроверкой с F1 fraud = {fmt(to_float(exp2_qwen[('llm_self_check', 'subtle')]['f1_fraud']))}.",
    )
    add_paragraph(
        document,
        "Эти результаты подтверждают, что LLM опираются не только на фиксированные словесные паттерны, но и на более высокий уровень "
        "семантического понимания. Особенно важно, что даже в условиях ASR-подобного шума и частично завуалированных инструкций "
        "модели сохраняют рабочее качество, тогда как правила почти полностью теряют чувствительность к атаке.",
    )
    add_figure(
        document,
        OUTPUT_DIR / "report_charts" / "exp2_f1_heatmap.png",
        "Рисунок 2 – Тепловая карта F1 fraud для различных архитектур и вариантов аугментации",
        width_cm=15.0,
    )
    add_figure(
        document,
        OUTPUT_DIR / "report_charts" / "exp2_robustness_lines.png",
        "Рисунок 3 – Сравнение устойчивости базового классификатора на основе правил и LLM-конфигураций к аугментациям",
        width_cm=15.0,
    )

    add_heading(document, "1.7 Результаты внешнего бенчмарка", 2)
    exp3_rows = [
        ["Правила", fmt(to_float(exp3_rules["accuracy"])), fmt(to_float(exp3_rules["precision_fraud"])), fmt(to_float(exp3_rules["recall_fraud"])), fmt(to_float(exp3_rules["f1_fraud"])), exp3_rules["false_positives"], exp3_rules["false_negatives"]],
        ["Qwen 2.5-14B: прямая классификация", fmt(to_float(exp3_qwen["single_llm"]["accuracy"])), fmt(to_float(exp3_qwen["single_llm"]["precision_fraud"])), fmt(to_float(exp3_qwen["single_llm"]["recall_fraud"])), fmt(to_float(exp3_qwen["single_llm"]["f1_fraud"])), exp3_qwen["single_llm"]["false_positives"], exp3_qwen["single_llm"]["false_negatives"]],
        ["Qwen 2.5-14B: чек-лист признаков", fmt(to_float(exp3_qwen["llm_checklist"]["accuracy"])), fmt(to_float(exp3_qwen["llm_checklist"]["precision_fraud"])), fmt(to_float(exp3_qwen["llm_checklist"]["recall_fraud"])), fmt(to_float(exp3_qwen["llm_checklist"]["f1_fraud"])), exp3_qwen["llm_checklist"]["false_positives"], exp3_qwen["llm_checklist"]["false_negatives"]],
        ["Qwen 2.5-14B: самопроверка", fmt(to_float(exp3_qwen["llm_self_check"]["accuracy"])), fmt(to_float(exp3_qwen["llm_self_check"]["precision_fraud"])), fmt(to_float(exp3_qwen["llm_self_check"]["recall_fraud"])), fmt(to_float(exp3_qwen["llm_self_check"]["f1_fraud"])), exp3_qwen["llm_self_check"]["false_positives"], exp3_qwen["llm_self_check"]["false_negatives"]],
        ["Qwen 2.5-14B: ансамбль", fmt(to_float(exp3_qwen["llm_ensemble"]["accuracy"])), fmt(to_float(exp3_qwen["llm_ensemble"]["precision_fraud"])), fmt(to_float(exp3_qwen["llm_ensemble"]["recall_fraud"])), fmt(to_float(exp3_qwen["llm_ensemble"]["f1_fraud"])), exp3_qwen["llm_ensemble"]["false_positives"], exp3_qwen["llm_ensemble"]["false_negatives"]],
    ]
    add_table(
        document,
        ["Подход", "Accuracy", "Precision", "Recall", "F1", "FP", "FN"],
        exp3_rows,
        [5.7, 1.7, 1.7, 1.7, 1.4, 1.2, 1.2],
    )
    add_caption(document, "Таблица 5 – Результаты Эксперимента 3 на внешнем бенчмарке")
    add_paragraph(
        document,
        f"На внешнем бенчмарке базовый классификатор на основе правил оказался чрезмерно консервативным: при precision fraud = "
        f"{fmt(to_float(exp3_rules['precision_fraud']))} его recall fraud составил лишь {fmt(to_float(exp3_rules['recall_fraud']))}, "
        f"а F1 fraud – {fmt(to_float(exp3_rules['f1_fraud']))}. Все LLM-конфигурации многократно превзошли этот подход. "
        f"Максимальная accuracy была достигнута конфигурацией с самопроверкой и составила {fmt(to_float(exp3_qwen['llm_self_check']['accuracy']))}. "
        f"Лучшее значение F1 fraud ({fmt(to_float(exp3_qwen['single_llm']['f1_fraud']))}) одновременно показали прямая LLM-классификация, "
        "конфигурация с самопроверкой и ансамблевая конфигурация.",
    )
    add_paragraph(
        document,
        "Наблюдаемая картина свидетельствует о хорошей переносимости разработанной постановки на данные, не совпадающие с базовым "
        "синтетическим корпусом. Иначе говоря, модели научились не просто запоминать внутренние шаблоны датасета, а распознавать "
        "общие признаки социальной инженерии и опасных инструкций. Это особенно важно для прикладного использования детектора в системах мониторинга.",
    )
    add_figure(
        document,
        OUTPUT_DIR / "report_charts" / "exp3_f1_comparison.png",
        "Рисунок 4 – Сравнение базового классификатора на основе правил и LLM-конфигураций по F1 fraud на внешнем бенчмарке",
    )

    add_heading(document, "1.8 Дополнительные исследования", 2)
    version_rows = []
    for architecture in ["single_llm", "llm_checklist", "llm_self_check", "llm_ensemble"]:
        old_f1 = to_float(exp1_old[architecture]["f1_fraud"])
        new_f1 = to_float(exp1_qwen[architecture]["f1_fraud"])
        version_rows.append(
            [
                ARCHITECTURE_NAMES_RU[architecture],
                fmt(old_f1),
                fmt(new_f1),
                f"+{fmt(new_f1 - old_f1)}" if new_f1 >= old_f1 else fmt(new_f1 - old_f1),
            ]
        )
    add_table(
        document,
        ["Архитектура", "F1 v1", "F1 v2", "Изменение"],
        version_rows,
        [5.0, 3.0, 3.0, 3.0],
    )
    add_caption(document, "Таблица 6 – Сравнение двух версий Эксперимента 1")
    add_paragraph(
        document,
        "Сравнение двух версий первого эксперимента показало, что переработка промптовой конфигурации практически не изменила качество "
        "прямой LLM-классификации, но радикально улучшила структурированные режимы. Наибольший прирост показала конфигурация с самопроверкой: "
        f"ее F1 fraud увеличился с {fmt(to_float(exp1_old['llm_self_check']['f1_fraud']))} до "
        f"{fmt(to_float(exp1_qwen['llm_self_check']['f1_fraud']))}. Это означает, что для сложных промптовых архитектур качество "
        "формулировки инструкции и порядок самоанализа модели оказывают критическое влияние на результат.",
    )
    add_figure(
        document,
        OUTPUT_DIR / "report_charts" / "supplement_exp1_versions.png",
        "Рисунок 5 – Влияние обновления промптовой конфигурации на результаты Эксперимента 1",
        width_cm=14.5,
    )

    gemma_rows = [
        ["Правила", "набор правил", fmt(to_float(exp1_rules["f1_fraud"])), fmt(to_float(exp1_rules["recall_fraud"])), fmt(to_float(exp1_rules["latency_ms_avg"]))],
        ["Gemma 4-31B API", "прямая классификация", fmt(to_float(gemma_exp1["f1_fraud"])), fmt(to_float(gemma_exp1["recall_fraud"])), fmt(to_float(gemma_exp1["latency_ms_avg"]))],
        ["Qwen 2.5-14B", "самопроверка", fmt(to_float(exp1_qwen["llm_self_check"]["f1_fraud"])), fmt(to_float(exp1_qwen["llm_self_check"]["recall_fraud"])), fmt(to_float(exp1_qwen["llm_self_check"]["latency_ms_avg"]))],
    ]
    add_table(
        document,
        ["Подход", "Конфигурация", "F1 fraud", "Recall fraud", "Latency, мс"],
        gemma_rows,
        [4.3, 3.2, 2.2, 2.5, 3.0],
    )
    add_caption(document, "Таблица 7 – Дополнительное сравнение Qwen и Gemma на Эксперименте 1")
    add_paragraph(
        document,
        f"В дополнительном API-эксперименте модель Gemma 4-31B в конфигурации прямой классификации показала F1 fraud = "
        f"{fmt(to_float(gemma_exp1['f1_fraud']))}, что немного выше результата базового классификатора на основе правил "
        f"({fmt(to_float(exp1_rules['f1_fraud']))}), но ниже выбранной рабочей конфигурации Qwen "
        f"с самопроверкой ({fmt(to_float(exp1_qwen['llm_self_check']['f1_fraud']))}). "
        f"При этом средняя задержка Gemma составила {fmt(to_float(gemma_exp1['latency_ms_avg']))} мс на пример, "
        "что на порядок выше правилового решения и заметно выше локального Qwen-пайплайна.",
    )
    add_paragraph(
        document,
        "Интерпретация этого результата должна быть аккуратной. Для Gemma в рамках работы был завершен только прогон в конфигурации прямой классификации, "
        "тогда как для Qwen исследовались несколько специализированных промптовых схем. Следовательно, полученное сравнение следует "
        "рассматривать как прикладное дополнительное исследование, а не как полностью симметричный модельный бенчмарк.",
    )
    add_figure(
        document,
        OUTPUT_DIR / "report_charts" / "supplement_exp4_exp1_gemma.png",
        "Рисунок 6 – Дополнительное сравнение базового классификатора на основе правил, Qwen и Gemma на Эксперименте 1",
        width_cm=14.5,
    )

    add_heading(document, "1.9 Анализ полученных результатов", 2)
    add_paragraph(
        document,
        "Анализ результатов всех проведенных экспериментов позволяет заключить, что наиболее устойчивой и практически значимой "
        "конфигурацией в рассматриваемой постановке является конфигурация с самопроверкой. Несмотря на то, что прямая LLM-классификация "
        "продемонстрировала наивысшие значения F1 fraud на базовом наборе и на суммарных аугментированных вариантах, именно "
        "конфигурация с самопроверкой обеспечила наилучшую полноту на Эксперименте 1, лучший результат на наиболее сложном варианте "
        "subtle в Эксперименте 2, а также наивысшую accuracy на внешнем бенчмарке Эксперимента 3. Это позволяет сделать вывод о том, "
        "что механизм самопроверки повышает устойчивость модели к неоднозначным, завуалированным и пограничным случаям мошеннической коммуникации.",
    )
    add_paragraph(
        document,
        "С практической точки зрения конфигурация с самопроверкой представляет наибольший интерес в задачах, где критически важно "
        "минимизировать пропуск опасных сценариев. Самопроверка оказывается особенно полезной в ситуациях, когда мошеннические признаки "
        "выражены неявно, а итоговое решение требует более осторожного разграничения между действительно опасным воздействием и внешне "
        "похожим, но нейтральным сервисным или бытовым контекстом. В свою очередь, прямую LLM-классификацию можно рассматривать как "
        "эффективный базовый режим применения модели, обеспечивающий высокий уровень итогового качества при более простой и менее затратной схеме использования.",
    )
    add_paragraph(
        document,
        "Базовый классификатор на основе правил, напротив, продемонстрировал две принципиальные слабости. Во-первых, он систематически "
        "уступал LLM-конфигурациям по полноте, то есть чаще допускал пропуск реальных мошеннических случаев. Во-вторых, его качество "
        "существенно снижалось при аугментациях, особенно при ASR-подобных искажениях текста. Данный результат подтверждает, что ручные "
        "сигнальные шаблоны плохо переносятся на естественные вариации речи, формулировок и степени явности мошеннических признаков. "
        "Вместе с тем такой подход сохраняет прикладную ценность как быстрый, интерпретируемый и вычислительно дешевый первый фильтр для наиболее очевидных сценариев.",
    )
    add_paragraph(
        document,
        "Отдельного внимания заслуживает компромисс между качеством классификации и вычислительной стоимостью. Базовый классификатор "
        f"на основе правил работает практически мгновенно, в среднем около {fmt(to_float(exp1_rules['latency_ms_avg']))} мс на один пример, "
        "тогда как LLM-подходы требуют существенно большего времени, особенно в случае использования внешнего API. В связи с этим "
        "практическая архитектура реальной системы может быть организована по каскадному принципу: на первом этапе выполняется быстрое "
        "отсечение наиболее очевидных случаев с помощью правил, а на втором этапе проводится более глубокий LLM-анализ спорных, "
        "неоднозначных или наиболее рискованных диалогов.",
    )
    add_figure(
        document,
        OUTPUT_DIR / "report_charts" / "overall_best_models.png",
        "Рисунок 7 – Сводное сравнение базового классификатора на основе правил и лучшей LLM-конфигурации по основным экспериментам",
        width_cm=14.5,
    )

    add_heading(document, "1.10 Выводы по выполненному заданию", 2)
    add_paragraph(
        document,
        "В ходе выполнения индивидуального задания был разработан воспроизводимый экспериментальный стенд для исследования методов "
        "выявления мошеннических диалогов на русском языке. Сформированы сценарный каталог и словарь канонических сигналов, подготовлены "
        "базовый и аугментированные наборы данных, разработаны промптовые конфигурации, правила классификации и скрипты автоматического "
        "расчета метрик и построения графиков.",
    )
    add_paragraph(
        document,
        "Полученные результаты показывают, что подходы, основанные на использовании больших языковых моделей, уверенно превосходят "
        "базовый классификатор на основе правил по качеству на основном наборе, сохраняют преимущество на аугментированных данных "
        "и демонстрируют хорошую переносимость на внешний бенчмарк. Наиболее эффективной по совокупности агрегированных показателей "
        "оказалась конфигурация прямой LLM-классификации, тогда как конфигурация с самопроверкой продемонстрировала наибольшую "
        "устойчивость при анализе завуалированных и пограничных случаев. Дополнительные эксперименты подтвердили высокую чувствительность "
        "итогового качества к формулировке промпта и показали, что облачная модель Gemma в простой конфигурации превосходит базовый "
        "классификатор на основе правил, однако уступает выбранной рабочей конфигурации Qwen.",
    )

    add_heading(document, "ЗАКЛЮЧЕНИЕ", 1)
    add_paragraph(
        document,
        "В течение практики по научно-исследовательской работе поставленные задачи были успешно выполнены. В ходе работы был проведен "
        "анализ предметной области выявления мошеннических диалогов, обоснована актуальность применения методов автоматической "
        "классификации в задачах защиты пользователей от социальной инженерии, а также сформулированы цель, объект, предмет и "
        "программа исследования. Кроме того, был разработан воспроизводимый экспериментальный стенд, включающий каталог сценариев, "
        "систему канонических сигналов мошенничества, набор промптовых конфигураций, процедуры оценки качества и средства визуализации результатов.",
    )
    add_paragraph(
        document,
        "Основной результат исследования состоит в том, что подходы, основанные на использовании больших языковых моделей, существенно "
        "превосходят базовый классификатор на основе правил как на основном синтетическом наборе, так и на аугментированных данных "
        "и внешнем бенчмарке. Наиболее эффективной по совокупности агрегированных показателей оказалась конфигурация прямой "
        "LLM-классификации, тогда как конфигурация с самопроверкой продемонстрировала наилучшую полноту и наибольшую устойчивость "
        "при анализе завуалированных сценариев. Отдельно установлено, что обновление промптовой конфигурации существенно улучшает "
        "качество структурированных режимов классификации, а дополнительный API-прогон модели Gemma показывает положительный, "
        "хотя и более скромный результат по сравнению с рабочей конфигурацией Qwen.",
    )
    add_paragraph(
        document,
        "За время прохождения практики освоены необходимые компетенции в области разработки и экспериментальной оценки методов защиты "
        "информации, подготовки и анализа датасетов, применения LLM в прикладных задачах и интерпретации полученных результатов. "
        "Полученные материалы могут быть использованы как основа для дальнейшего развития детектора мошеннических диалогов, расширения "
        "корпуса данных и разработки каскадной схемы практического применения в системах мониторинга и пользовательской защиты.",
    )

    add_heading(document, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1)
    sources = [
        "Обзор операций, совершенных без добровольного согласия клиентов финансовых организаций // Банк России, 2025. URL: https://www.cbr.ru/analytics/ib/operations_survey/2025/",
        "Qwen2.5 Technical Report // arXiv:2412.15115, 2024. URL: https://arxiv.org/abs/2412.15115",
        "Generating content // Gemini API, Google AI for Developers. URL: https://ai.google.dev/api/generate-content",
        "Li H., Huang S., Park J. K. A Human-Centered Review of Large Language Models for Online Scam Detection // SSRN, 2026. URL: https://papers.ssrn.com/sol3/papers.cfm?abstract_id=6338940",
        "SCRIPTMIND: Crime Script Inference and Cognitive Evaluation for LLM-based Social Engineering Scam Detection System // arXiv:2601.13581, 2026. URL: https://arxiv.org/abs/2601.13581",
        "PreScam: A Benchmark for Predicting Scam Progression from Early Conversations // arXiv:2605.12243, 2026. URL: https://arxiv.org/abs/2605.12243",
        "Wang X. et al. Self-Consistency Improves Chain of Thought Reasoning in Language Models // ICLR, 2023. URL: https://arxiv.org/abs/2203.11171",
        "Wei J., Zou K. EDA: Easy Data Augmentation Techniques for Boosting Performance on Text Classification Tasks // EMNLP-IJCNLP, 2019. URL: https://arxiv.org/abs/1901.11196",
        "Minaee S. et al. Large Language Models: A Survey // arXiv:2402.06196, 2024. URL: https://arxiv.org/abs/2402.06196",
        "Промптинг, аугментация и валидация в задачах классификации русскоязычных мошеннических диалогов: внутренние экспериментальные материалы проекта, 2026.",
    ]
    for index, source in enumerate(sources, start=1):
        add_paragraph(document, f"{index}. {source}", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, size=12)

    add_page_break(document)
    add_heading(document, "ОТЗЫВ О ПРОХОЖДЕНИИ ПРАКТИКИ", 1)
    add_paragraph(document, "Вид практики: производственная практика", first_line_indent=None)
    add_paragraph(document, "Тип практики: научно-исследовательская работа", first_line_indent=None)
    add_paragraph(
        document,
        f"Сроки прохождения практики: с {PRACTICE_START} г. по {PRACTICE_END} г.",
        first_line_indent=None,
    )
    add_paragraph(
        document,
        "по направлению подготовки 10.05.03 Информационная безопасность автоматизированных систем",
        first_line_indent=None,
    )
    add_paragraph(
        document,
        "направленность (профиль) «Безопасность открытых информационных систем»",
        first_line_indent=None,
    )
    add_paragraph(document, f"студентом группы № {GROUP} {STUDENT_NAME}", first_line_indent=None)
    add_table(
        document,
        ["№ п/п", "Критерии оценки", "Оценка (по 5-балльной шкале)"],
        [
            ["1", "Общая систематичность и ответственность работы в ходе практики", ""],
            ["2", "Достижение планируемых результатов практики", ""],
            ["3", "Корректность в сборе, анализе и интерпретации представляемых данных", ""],
            ["4", "Степень личного участия и самостоятельности практиканта в представляемом отчете о практике", ""],
            ["5", "Качество оформления отчетной документации", ""],
        ],
        [1.2, 11.8, 4.0],
    )
    add_paragraph(document, "ИТОГОВАЯ ОЦЕНКА ____________________", first_line_indent=None)
    add_paragraph(document, "", first_line_indent=None)
    add_paragraph(document, f"Руководитель практики от университета: {SUPERVISOR}", first_line_indent=None)
    add_paragraph(document, "(подпись) ____________________", first_line_indent=None)
    add_paragraph(
        document,
        "* Итоговая оценка выставляется как средняя арифметическая оценок по пяти критериям оценки.",
        first_line_indent=None,
        size=12,
    )

    document.save(REPORT_PATH)
    return REPORT_PATH


def main() -> None:
    path = build_report()
    print(f"Wrote NIRS report draft to {path}")


if __name__ == "__main__":
    main()
